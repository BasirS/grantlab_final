"""
Document management routes for uploading, processing, and searching grant documents.
"""

from typing import List, Optional
import os
from pathlib import Path

from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func
from docx import Document as DocxDocument
import PyPDF2
import structlog

from src.core.database import get_db, User, Document
from src.core.vector_store import get_vector_store
from src.api.routes.auth import get_current_user
from src.api.models.schemas import (
    DocumentCreate,
    DocumentResponse,
    DocumentSearchRequest,
    DocumentSearchResult,
    DocumentBatchProcess,
    DocumentBatchResult,
    DocumentUpload,
)
from src.config import settings


# initializing structured logger
logger = structlog.get_logger(__name__)

# creating router for document endpoints
router = APIRouter(prefix="/documents", tags=["Documents"])


async def extract_text_from_file(file_path: str, mime_type: str) -> str:
    """
    extracting text content from uploaded file based on mime type
    """
    try:
        if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_path.endswith(".docx"):
            # extracting text from docx file
            doc = DocxDocument(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text

        elif mime_type == "application/pdf" or file_path.endswith(".pdf"):
            # extracting text from pdf file
            with open(file_path, "rb") as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text

        elif mime_type == "text/plain" or file_path.endswith(".txt"):
            # reading plain text file
            with open(file_path, "r", encoding="utf-8", errors="ignore") as txt_file:
                return txt_file.read()

        else:
            raise ValueError(f"Unsupported file type: {mime_type}")

    except Exception as e:
        logger.error("text_extraction_error", file_path=file_path, error=str(e))
        raise


@router.post("/upload", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def upload_document(
    file: UploadFile = File(...),
    document_type: str = Form(...),
    program: Optional[str] = Form(None),
    year: Optional[int] = Form(None),
    funder: Optional[str] = Form(None),
    funder_type: Optional[str] = Form(None),
    outcome: Optional[str] = Form(None),
    award_amount: Optional[float] = Form(None),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    uploading document file and extracting content for processing
    """
    try:
        logger.info(
            "uploading_document",
            filename=file.filename,
            document_type=document_type,
            user_id=current_user.id,
        )

        # validating file extension
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in settings.allowed_file_extensions:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File type {file_ext} not allowed. Allowed types: {settings.allowed_file_extensions}",
            )

        # validating file size
        file.file.seek(0, 2)  # seeking to end of file
        file_size = file.file.tell()
        file.file.seek(0)  # seeking back to start

        if file_size > settings.max_upload_size_bytes:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File size {file_size} bytes exceeds maximum {settings.max_upload_size_bytes} bytes",
            )

        # creating upload directory if it doesn't exist
        upload_dir = Path(settings.upload_directory)
        upload_dir.mkdir(parents=True, exist_ok=True)

        # saving file to disk
        file_path = upload_dir / f"{current_user.id}_{file.filename}"
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)

        # extracting text content
        try:
            text_content = await extract_text_from_file(str(file_path), file.content_type)
        except Exception as e:
            # removing file if extraction fails
            os.remove(file_path)
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Failed to extract text from file: {str(e)}",
            )

        # creating document record
        document = Document(
            filename=file.filename,
            document_type=document_type,
            program=program,
            year=year,
            funder=funder,
            funder_type=funder_type,
            outcome=outcome,
            award_amount=award_amount,
            content=text_content,
            file_path=str(file_path),
            file_size_bytes=file_size,
            mime_type=file.content_type,
            processed=False,
        )

        db.add(document)
        await db.commit()
        await db.refresh(document)

        # generating and storing embedding asynchronously
        try:
            vector_store = get_vector_store(db)
            await vector_store.add_document_embedding(document.id, text_content)
            logger.info(
                "document_uploaded_and_processed",
                document_id=document.id,
                filename=file.filename,
            )
        except Exception as e:
            logger.error(
                "embedding_generation_failed",
                document_id=document.id,
                error=str(e),
            )
            # updating document with error but keeping the document record
            document.processing_error = str(e)
            await db.commit()

        return DocumentResponse.from_orm(document)

    except HTTPException:
        raise
    except Exception as e:
        logger.error("upload_error", filename=file.filename, error=str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to upload document: {str(e)}",
        )


@router.get("/", response_model=List[DocumentResponse])
async def list_documents(
    document_type: Optional[str] = None,
    program: Optional[str] = None,
    funder: Optional[str] = None,
    processed: Optional[bool] = None,
    limit: int = 50,
    offset: int = 0,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    listing documents with optional filters
    """
    try:
        logger.debug(
            "listing_documents",
            document_type=document_type,
            program=program,
            limit=limit,
            offset=offset,
        )

        # building query with filters
        query = select(Document)

        if document_type:
            query = query.where(Document.document_type == document_type)
        if program:
            query = query.where(Document.program == program)
        if funder:
            query = query.where(Document.funder == funder)
        if processed is not None:
            query = query.where(Document.processed == processed)

        query = query.order_by(Document.created_at.desc()).limit(limit).offset(offset)

        result = await db.execute(query)
        documents = result.scalars().all()

        logger.info("documents_listed", count=len(documents))

        return [DocumentResponse.from_orm(doc) for doc in documents]

    except Exception as e:
        logger.error("list_error", error=str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to list documents",
        )


@router.get("/{document_id}", response_model=DocumentResponse)
async def get_document(
    document_id: int,
    include_content: bool = False,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    getting specific document by id with optional content inclusion
    """
    try:
        logger.debug("fetching_document", document_id=document_id)

        result = await db.execute(
            select(Document).where(Document.id == document_id)
        )
        document = result.scalar_one_or_none()

        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found",
            )

        doc_response = DocumentResponse.from_orm(document)

        if include_content:
            doc_response.content = document.content

        return doc_response

    except HTTPException:
        raise
    except Exception as e:
        logger.error("get_document_error", document_id=document_id, error=str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to get document",
        )


@router.post("/search", response_model=List[DocumentSearchResult])
async def search_documents(
    search_request: DocumentSearchRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    searching documents using semantic similarity
    """
    try:
        logger.info(
            "searching_documents",
            query_length=len(search_request.query),
            document_type=search_request.document_type,
            top_k=search_request.top_k,
        )

        vector_store = get_vector_store(db)

        results = await vector_store.search_similar_documents(
            query_text=search_request.query,
            document_type=search_request.document_type,
            program=search_request.program,
            funder=search_request.funder,
            top_k=search_request.top_k,
            similarity_threshold=search_request.similarity_threshold,
        )

        search_results = [
            DocumentSearchResult(
                document=DocumentResponse.from_orm(doc),
                similarity_score=score,
            )
            for doc, score in results
        ]

        logger.info("search_complete", results_count=len(search_results))

        return search_results

    except Exception as e:
        logger.error("search_error", error=str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to search documents: {str(e)}",
        )


@router.post("/batch-process", response_model=DocumentBatchResult)
async def batch_process_documents(
    batch_request: DocumentBatchProcess,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    processing multiple documents to generate embeddings in batch
    """
    try:
        logger.info(
            "batch_processing",
            document_count=len(batch_request.document_ids),
            user_id=current_user.id,
        )

        vector_store = get_vector_store(db)

        results = await vector_store.batch_add_document_embeddings(
            document_ids=batch_request.document_ids,
        )

        logger.info(
            "batch_processing_complete",
            success=len(results["success"]),
            failed=len(results["failed"]),
        )

        return DocumentBatchResult(**results)

    except Exception as e:
        logger.error("batch_process_error", error=str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to batch process documents: {str(e)}",
        )


@router.delete("/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    deleting document and its associated file
    """
    try:
        logger.info("deleting_document", document_id=document_id, user_id=current_user.id)

        result = await db.execute(
            select(Document).where(Document.id == document_id)
        )
        document = result.scalar_one_or_none()

        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found",
            )

        # deleting file from disk if it exists
        if document.file_path and os.path.exists(document.file_path):
            try:
                os.remove(document.file_path)
                logger.debug("file_deleted", file_path=document.file_path)
            except Exception as e:
                logger.warning("file_deletion_failed", file_path=document.file_path, error=str(e))

        # deleting document from database
        await db.delete(document)
        await db.commit()

        logger.info("document_deleted", document_id=document_id)

        return None

    except HTTPException:
        raise
    except Exception as e:
        logger.error("delete_error", document_id=document_id, error=str(e))
        await db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to delete document",
        )


@router.get("/stats/summary")
async def get_document_stats(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    getting summary statistics about documents
    """
    try:
        logger.debug("fetching_document_stats")

        # getting total count
        total_result = await db.execute(select(func.count(Document.id)))
        total_count = total_result.scalar()

        # getting count by document type
        type_result = await db.execute(
            select(Document.document_type, func.count(Document.id))
            .group_by(Document.document_type)
        )
        type_breakdown = {row[0]: row[1] for row in type_result.all()}

        # getting count by program
        program_result = await db.execute(
            select(Document.program, func.count(Document.id))
            .where(Document.program.isnot(None))
            .group_by(Document.program)
        )
        program_breakdown = {row[0]: row[1] for row in program_result.all()}

        # getting processed vs unprocessed count
        processed_result = await db.execute(
            select(Document.processed, func.count(Document.id))
            .group_by(Document.processed)
        )
        processed_breakdown = {row[0]: row[1] for row in processed_result.all()}

        stats = {
            "total_documents": total_count,
            "by_type": type_breakdown,
            "by_program": program_breakdown,
            "processed": processed_breakdown.get(True, 0),
            "unprocessed": processed_breakdown.get(False, 0),
        }

        logger.info("document_stats_fetched", total=total_count)

        return stats

    except Exception as e:
        logger.error("stats_error", error=str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to get document statistics",
        )
