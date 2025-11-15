"""
Professional Word document export service for grant applications.

This module provides functionality to generate professionally formatted Word documents
matching grant application standards with proper section headers, Cambio Labs branding,
and formatting that aligns with reference grant documents.
"""

import os
import tempfile
from datetime import datetime
from typing import Dict, Any, List, Optional
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import structlog

from src.core.database import GrantApplication, User


# initializing structured logger
logger = structlog.get_logger(__name__)


class WordDocumentExporter:
    """
    service for exporting grant applications to professionally formatted word documents

    handles document creation with proper styling, branding, and formatting that matches
    grant application standards including cambio labs letterhead, section hierarchy,
    and professional layout suitable for submission
    """

    # cambio labs branding colors
    CAMBIO_PURPLE = RGBColor(106, 27, 154)  # primary purple
    CAMBIO_DARK_BLUE = RGBColor(0, 0, 139)  # dark blue for headings
    TEXT_COLOR = RGBColor(0, 0, 0)  # black for body text
    HEADER_BG = RGBColor(245, 245, 245)  # light gray for header background

    # standard grant sections in order
    SECTION_ORDER = [
        ("executive_summary", "Executive Summary"),
        ("organizational_background", "Organizational Background"),
        ("problem_statement", "Need Statement"),
        ("project_description", "Project Description"),
        ("budget_narrative", "Budget Narrative"),
        ("evaluation_plan", "Evaluation Plan"),
    ]

    def __init__(self, logo_path: Optional[str] = None):
        """
        initializing word document exporter

        args:
            logo_path: path to cambio labs logo image file (optional)
        """
        self.logo_path = logo_path

        # checking if logo exists
        if self.logo_path and not os.path.exists(self.logo_path):
            logger.warning("logo_not_found", path=self.logo_path)
            self.logo_path = None

    async def create_grant_document(
        self,
        grant: GrantApplication,
        user: User,
        output_path: Optional[str] = None,
        include_metadata: bool = True,
    ) -> str:
        """
        creating professionally formatted word document for grant application

        args:
            grant: grant application model from database
            user: user who created the grant
            output_path: custom output path (optional, defaults to temp file)
            include_metadata: whether to include generation metadata

        returns:
            path to generated word document
        """
        try:
            # initializing new word document
            doc = Document()

            # setting up document margins and page setup
            self._setup_page_layout(doc)

            # creating custom styles
            self._create_custom_styles(doc)

            # adding letterhead with cambio labs branding
            self._add_letterhead(doc, user)

            # adding cover page with grant metadata
            self._add_cover_page(doc, grant, user)

            # adding page break after cover
            doc.add_page_break()

            # adding table of contents
            self._add_table_of_contents(doc, grant)

            # adding page break before main content
            doc.add_page_break()

            # adding main grant sections
            self._add_grant_sections(doc, grant)

            # adding metadata footer if requested
            if include_metadata:
                self._add_metadata_footer(doc, grant)

            # determining output path
            if not output_path:
                # creating temporary file
                temp_dir = tempfile.gettempdir()
                safe_title = "".join(c for c in grant.title if c.isalnum() or c in (' ', '-', '_'))
                safe_title = safe_title.replace(' ', '_')[:50]  # limiting filename length
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_path = os.path.join(
                    temp_dir,
                    f"cambio_labs_{safe_title}_{timestamp}.docx"
                )

            # saving document
            doc.save(output_path)

            logger.info(
                "word_document_created",
                grant_id=grant.id,
                output_path=output_path,
                file_size=os.path.getsize(output_path)
            )

            return output_path

        except Exception as e:
            logger.error(
                "word_export_error",
                grant_id=grant.id,
                error=str(e),
                exc_info=True
            )
            raise

    def _setup_page_layout(self, doc: Document):
        """
        setting up page margins and layout

        standard 1-inch margins on all sides for professional grant formatting
        """
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.page_height = Inches(11)  # letter size
            section.page_width = Inches(8.5)

    def _create_custom_styles(self, doc: Document):
        """
        creating custom styles for consistent formatting throughout document
        """
        styles = doc.styles

        # creating custom heading 1 style
        if 'Cambio Heading 1' not in styles:
            heading1_style = styles.add_style('Cambio Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            heading1_font = heading1_style.font
            heading1_font.name = 'Calibri'
            heading1_font.size = Pt(14)
            heading1_font.bold = True
            heading1_font.color.rgb = self.CAMBIO_DARK_BLUE
            heading1_style.paragraph_format.space_before = Pt(12)
            heading1_style.paragraph_format.space_after = Pt(6)
            heading1_style.paragraph_format.keep_with_next = True

        # creating custom heading 2 style
        if 'Cambio Heading 2' not in styles:
            heading2_style = styles.add_style('Cambio Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            heading2_font = heading2_style.font
            heading2_font.name = 'Calibri'
            heading2_font.size = Pt(12)
            heading2_font.bold = True
            heading2_font.color.rgb = self.CAMBIO_DARK_BLUE
            heading2_style.paragraph_format.space_before = Pt(10)
            heading2_style.paragraph_format.space_after = Pt(4)

        # creating custom body text style
        if 'Cambio Body' not in styles:
            body_style = styles.add_style('Cambio Body', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = 'Calibri'
            body_font.size = Pt(11)
            body_font.color.rgb = self.TEXT_COLOR
            body_style.paragraph_format.space_after = Pt(12)
            body_style.paragraph_format.line_spacing = 1.15
            body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _add_letterhead(self, doc: Document, user: User):
        """
        adding professional letterhead with cambio labs branding

        includes logo (if available) and organization contact information
        """
        # creating header section
        section = doc.sections[0]
        header = section.header

        # clearing any existing header content
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

        # creating table for letterhead layout
        table = header.add_table(rows=1, cols=2, width=Inches(6.5))
        table.autofit = False

        # left cell for logo or organization name
        left_cell = table.rows[0].cells[0]
        left_cell.width = Inches(2.5)

        if self.logo_path:
            # adding logo image
            logo_para = left_cell.paragraphs[0]
            logo_run = logo_para.add_run()
            logo_run.add_picture(self.logo_path, width=Inches(1.5))
        else:
            # adding text-based branding
            org_para = left_cell.paragraphs[0]
            org_run = org_para.add_run("CAMBIO LABS")
            org_run.font.name = 'Calibri'
            org_run.font.size = Pt(16)
            org_run.font.bold = True
            org_run.font.color.rgb = self.CAMBIO_PURPLE

        # right cell for contact information
        right_cell = table.rows[0].cells[1]
        right_cell.width = Inches(4.0)

        contact_para = right_cell.paragraphs[0]
        contact_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # adding contact details
        contact_info = [
            "Cambio Labs",
            "Queens County, New York, United States",
            "(301) 717-9982",
            "sebastian@cambiolabs.org",
            "www.cambiolabs.org"
        ]

        for i, line in enumerate(contact_info):
            if i > 0:
                contact_para.add_run("\n")
            run = contact_para.add_run(line)
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            run.font.color.rgb = self.TEXT_COLOR

        # adding bottom border to header
        self._add_paragraph_border(header_para, bottom=True)

    def _add_cover_page(self, doc: Document, grant: GrantApplication, user: User):
        """
        adding professional cover page with grant metadata
        """
        # adding main title
        title = doc.add_heading("GRANT PROPOSAL", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(20)
        title.runs[0].font.bold = True
        title.runs[0].font.color.rgb = self.CAMBIO_PURPLE

        # adding grant title
        grant_title = doc.add_heading(grant.title, level=1)
        grant_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        grant_title.runs[0].font.size = Pt(16)
        grant_title.runs[0].font.color.rgb = self.CAMBIO_DARK_BLUE

        # adding spacing
        doc.add_paragraph()

        # adding submission information section
        info_heading = doc.add_heading("Submission Information", level=2)
        info_heading.runs[0].font.color.rgb = self.CAMBIO_DARK_BLUE

        # creating info table
        info_table = doc.add_table(rows=0, cols=2)
        info_table.style = 'Light Grid Accent 1'

        # helper function to add info row
        def add_info_row(label: str, value: str):
            row = info_table.add_row()
            label_cell = row.cells[0]
            value_cell = row.cells[1]

            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.font.bold = True
            label_run.font.size = Pt(11)

            value_para = value_cell.paragraphs[0]
            value_run = value_para.add_run(value)
            value_run.font.size = Pt(11)

        # adding submission details
        add_info_row("Submission Date:", datetime.now().strftime("%B %d, %Y"))
        add_info_row("Organization:", "Cambio Labs")

        if grant.funder:
            add_info_row("Funder:", grant.funder)

        if grant.funder_agency:
            add_info_row("Agency:", grant.funder_agency)

        if grant.opportunity_number:
            add_info_row("Opportunity Number:", grant.opportunity_number)

        add_info_row("Status:", grant.status.title())
        add_info_row("Version:", str(grant.version))

        # adding contact information section
        doc.add_paragraph()
        contact_heading = doc.add_heading("Contact Information", level=2)
        contact_heading.runs[0].font.color.rgb = self.CAMBIO_DARK_BLUE

        contact_para = doc.add_paragraph()
        contact_para.add_run(f"Primary Contact: Sebastián Martín\n")
        contact_para.add_run(f"Title: Founder & CEO\n")
        contact_para.add_run(f"Email: sebastian@cambiolabs.org\n")
        contact_para.add_run(f"Phone: (301) 717-9982")

        for run in contact_para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

    def _add_table_of_contents(self, doc: Document, grant: GrantApplication):
        """
        adding table of contents with section list
        """
        # adding toc heading
        toc_heading = doc.add_heading("TABLE OF CONTENTS", level=1)
        toc_heading.runs[0].font.color.rgb = self.CAMBIO_DARK_BLUE
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # adding spacing
        doc.add_paragraph()

        # adding section list
        section_number = 1
        for section_key, section_title in self.SECTION_ORDER:
            # checking if section has content
            section_content = getattr(grant, section_key, None)
            if section_content and section_content.strip():
                toc_entry = doc.add_paragraph()
                toc_entry.paragraph_format.left_indent = Inches(0.5)

                entry_run = toc_entry.add_run(f"{section_number}. {section_title}")
                entry_run.font.name = 'Calibri'
                entry_run.font.size = Pt(12)
                entry_run.font.bold = True

                section_number += 1

    def _add_grant_sections(self, doc: Document, grant: GrantApplication):
        """
        adding main grant content sections with proper formatting
        """
        section_number = 1

        for section_key, section_title in self.SECTION_ORDER:
            # retrieving section content from grant model
            section_content = getattr(grant, section_key, None)

            # skipping empty sections
            if not section_content or not section_content.strip():
                logger.debug("skipping_empty_section", section=section_key)
                continue

            # adding section heading
            heading = doc.add_heading(f"{section_number}. {section_title.upper()}", level=1)
            heading.runs[0].font.size = Pt(14)
            heading.runs[0].font.color.rgb = self.CAMBIO_DARK_BLUE
            heading.runs[0].font.bold = True

            # splitting content into paragraphs
            paragraphs = section_content.split('\n\n')

            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph()

                    # checking for subsection headers (lines ending with colon)
                    if para_text.strip().endswith(':') and len(para_text) < 100:
                        # treating as subsection header
                        run = para.add_run(para_text.strip())
                        run.font.name = 'Calibri'
                        run.font.size = Pt(12)
                        run.font.bold = True
                        run.font.color.rgb = self.CAMBIO_DARK_BLUE
                        para.paragraph_format.space_before = Pt(10)
                        para.paragraph_format.space_after = Pt(6)
                    else:
                        # regular body text
                        run = para.add_run(para_text.strip())
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
                        run.font.color.rgb = self.TEXT_COLOR

                        para.paragraph_format.space_after = Pt(12)
                        para.paragraph_format.line_spacing = 1.15
                        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # adding spacing after section
            doc.add_paragraph()

            section_number += 1

    def _add_metadata_footer(self, doc: Document, grant: GrantApplication):
        """
        adding footer with generation metadata
        """
        section = doc.sections[0]
        footer = section.footer

        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # adding metadata text
        metadata_text = f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')} | "
        metadata_text += f"Grant ID: {grant.id} | Version: {grant.version}"

        metadata_run = footer_para.add_run(metadata_text)
        metadata_run.font.name = 'Calibri'
        metadata_run.font.size = Pt(8)
        metadata_run.font.color.rgb = RGBColor(128, 128, 128)  # gray
        metadata_run.font.italic = True

    def _add_paragraph_border(
        self,
        paragraph,
        top: bool = False,
        bottom: bool = False,
        left: bool = False,
        right: bool = False
    ):
        """
        adding border to paragraph element

        args:
            paragraph: docx paragraph object
            top: add top border
            bottom: add bottom border
            left: add left border
            right: add right border
        """
        p = paragraph._p if hasattr(paragraph, '_p') else paragraph._element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')

        def add_border_element(name: str):
            element = OxmlElement(f'w:{name}')
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '12')  # border size
            element.set(qn('w:space'), '1')
            element.set(qn('w:color'), '000000')  # black
            pBdr.append(element)

        if top:
            add_border_element('top')
        if bottom:
            add_border_element('bottom')
        if left:
            add_border_element('left')
        if right:
            add_border_element('right')

        pPr.insert_element_before(
            pBdr,
            'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku',
            'w:wordWrap', 'w:overflowPunct', 'w:topLinePunct',
            'w:autoSpaceDE', 'w:autoSpaceDN', 'w:bidi', 'w:adjustRightInd',
            'w:snapToGrid', 'w:spacing', 'w:ind', 'w:contextualSpacing',
            'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
            'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
            'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
            'w:pPrChange'
        )


# creating singleton instance
_exporter_instance = None


def get_word_exporter(logo_path: Optional[str] = None) -> WordDocumentExporter:
    """
    getting singleton instance of word document exporter

    args:
        logo_path: path to cambio labs logo (optional)

    returns:
        word document exporter instance
    """
    global _exporter_instance

    if _exporter_instance is None:
        _exporter_instance = WordDocumentExporter(logo_path=logo_path)

    return _exporter_instance
