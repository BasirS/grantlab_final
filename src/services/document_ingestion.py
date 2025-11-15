"""
Document ingestion service for loading, processing, and chunking grant documents.
"""

import os
import re
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
import asyncio

from docx import Document as DocxDocument
import PyPDF2
from dateutil import parser as date_parser
from langchain.text_splitter import RecursiveCharacterTextSplitter
import structlog

from src.config import settings


# initializing structured logger
logger = structlog.get_logger(__name__)


class DocumentLoader:
    """
    loading documents from various file formats
    """

    @staticmethod
    async def load_txt_file(file_path: str) -> str:
        """
        loading text file with encoding fallback strategy
        """
        encodings = ['utf-8', 'windows-1252', 'latin-1', 'ascii']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                logger.debug("txt_loaded", file_path=file_path, encoding=encoding)
                return content
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error("txt_load_error", file_path=file_path, error=str(e))
                raise

        # if all encodings fail, read as binary and decode with errors='ignore'
        with open(file_path, 'rb') as f:
            content = f.read().decode('utf-8', errors='ignore')
        logger.warning("txt_loaded_with_errors", file_path=file_path)
        return content

    @staticmethod
    async def load_docx_file(file_path: str) -> str:
        """
        loading docx file and extracting text while preserving structure
        """
        try:
            doc = DocxDocument(file_path)

            # extracting paragraphs with structure preservation
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # checking if paragraph is a heading
                    if para.style.name.startswith('Heading'):
                        paragraphs.append(f"\n## {text}\n")
                    else:
                        paragraphs.append(text)

            # extracting text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        paragraphs.append(row_text)

            content = '\n\n'.join(paragraphs)
            logger.debug("docx_loaded", file_path=file_path, paragraphs=len(paragraphs))
            return content

        except Exception as e:
            logger.error("docx_load_error", file_path=file_path, error=str(e))
            raise

    @staticmethod
    async def load_pdf_file(file_path: str) -> str:
        """
        loading pdf file and extracting text
        """
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                pages = []

                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        pages.append(f"[Page {page_num + 1}]\n{text}")

            content = '\n\n'.join(pages)
            logger.debug("pdf_loaded", file_path=file_path, pages=len(pages))
            return content

        except Exception as e:
            logger.error("pdf_load_error", file_path=file_path, error=str(e))
            raise

    @classmethod
    async def load_file(cls, file_path: str) -> Tuple[str, str]:
        """
        loading file based on extension and returning content with mime type
        """
        file_ext = os.path.splitext(file_path)[1].lower()

        if file_ext == '.txt':
            content = await cls.load_txt_file(file_path)
            return content, 'text/plain'
        elif file_ext == '.docx':
            content = await cls.load_docx_file(file_path)
            return content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif file_ext == '.pdf':
            content = await cls.load_pdf_file(file_path)
            return content, 'application/pdf'
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")


class MetadataExtractor:
    """
    extracting metadata from document content and filename
    """

    # patterns for detecting programs
    PROGRAM_PATTERNS = {
        'Journey Platform': ['journey', 'sparky', 'ai learning platform'],
        'StartUp NYCHA': ['startup nycha', 'nycha', 'public housing'],
        'RETI': ['reti', 'renewable energy', 'solar workforce'],
        'Cambio Solar': ['cambio solar', 'green jobs', 'solar training'],
        'Cambio Coding': ['cambio coding', 'coding bootcamp', 'programming'],
        'Social Entrepreneurship': ['social entrepreneurship', 'social enterprise'],
    }

    # patterns for detecting funder types
    FUNDER_TYPE_PATTERNS = {
        'federal': ['nsf', 'department of', 'federal', 'nih', 'doe', 'usda', 'epa'],
        'foundation': ['foundation', 'fund', 'trust', 'endowment'],
        'corporate': ['google', 'microsoft', 'amazon', 'apple', 'corporate'],
        'state': ['new york state', 'nyserda', 'state grant'],
        'local': ['nyc', 'new york city', 'city grant', 'council'],
    }

    # keywords indicating document types
    DOCUMENT_TYPE_KEYWORDS = {
        'rfp': ['request for proposal', 'rfp', 'solicitation', 'opportunity'],
        'org_info': ['about us', 'our mission', 'organization overview'],
        'past_grant': ['grant proposal', 'application', 'submitted'],
    }

    @staticmethod
    def extract_from_filename(filename: str) -> Dict[str, Any]:
        """
        extracting metadata from filename patterns
        """
        metadata = {}

        # removing extension
        name = os.path.splitext(filename)[0]

        # extracting date if present
        date_patterns = [
            r'\d{4}[-_]\d{2}[-_]\d{2}',  # 2023-01-15
            r'\d{2}[-_]\d{2}[-_]\d{4}',  # 01-15-2023
            r'\d{4}',                     # 2023
        ]

        for pattern in date_patterns:
            match = re.search(pattern, name)
            if match:
                try:
                    date_str = match.group(0)
                    parsed_date = date_parser.parse(date_str, fuzzy=True)
                    metadata['year'] = parsed_date.year
                    metadata['date'] = parsed_date.isoformat()
                    break
                except:
                    pass

        # extracting funder name from filename
        # pattern: "FunderName Grant" or "Grant for FunderName"
        funder_match = re.search(r'([A-Z][A-Za-z\s&]+?)(?:\s+Grant|\s+Proposal)', name)
        if funder_match:
            metadata['funder'] = funder_match.group(1).strip()

        # detecting revision or version info
        if 'rev' in name.lower() or 'version' in name.lower():
            metadata['is_revision'] = True

        logger.debug("filename_metadata_extracted", filename=filename, metadata=metadata)
        return metadata

    @classmethod
    def extract_from_content(cls, content: str) -> Dict[str, Any]:
        """
        extracting metadata from document content
        """
        metadata = {}
        content_lower = content.lower()

        # detecting programs mentioned
        programs = []
        for program, keywords in cls.PROGRAM_PATTERNS.items():
            for keyword in keywords:
                if keyword in content_lower:
                    programs.append(program)
                    break

        if programs:
            metadata['programs'] = list(set(programs))
            # using first program as primary
            metadata['program'] = programs[0]

        # detecting funder type
        for funder_type, keywords in cls.FUNDER_TYPE_PATTERNS.items():
            for keyword in keywords:
                if keyword in content_lower:
                    metadata['funder_type'] = funder_type
                    break
            if 'funder_type' in metadata:
                break

        # extracting mission statement if present
        mission_patterns = [
            r'(?:our mission is to|mission:|mission statement:)\s*(.{50,300}?)(?:\.|$)',
            r'(?:we are committed to|committed to)\s*(.{50,300}?)(?:\.|$)',
        ]

        for pattern in mission_patterns:
            match = re.search(pattern, content_lower, re.IGNORECASE)
            if match:
                metadata['mission_excerpt'] = match.group(1).strip()
                break

        # extracting award amount if mentioned
        amount_patterns = [
            r'\$\s*([\d,]+)\s*(?:award|granted|funding)',
            r'(?:award|grant)\s+amount:\s*\$\s*([\d,]+)',
        ]

        for pattern in amount_patterns:
            match = re.search(pattern, content_lower)
            if match:
                try:
                    amount_str = match.group(1).replace(',', '')
                    metadata['award_amount'] = float(amount_str)
                    break
                except:
                    pass

        # detecting outcome keywords
        if any(word in content_lower for word in ['awarded', 'congratulations', 'approved']):
            metadata['outcome'] = 'awarded'
        elif any(word in content_lower for word in ['declined', 'not selected', 'unsuccessful']):
            metadata['outcome'] = 'declined'

        # extracting team members mentioned
        team_pattern = r'(?:team member|staff|director|coordinator):\s*([A-Z][a-z]+\s+[A-Z][a-z]+)'
        team_matches = re.findall(team_pattern, content, re.IGNORECASE)
        if team_matches:
            metadata['team_members'] = list(set(team_matches))

        logger.debug("content_metadata_extracted", metadata_keys=list(metadata.keys()))
        return metadata

    @classmethod
    def infer_document_type(cls, filename: str, content: str, folder_name: str) -> str:
        """
        inferring document type from filename, content, and folder structure
        """
        filename_lower = filename.lower()
        content_lower = content.lower()
        folder_lower = folder_name.lower()

        # checking folder structure first
        if 'rfp' in folder_lower or 'solicitation' in folder_lower:
            return 'rfp'
        elif 'org' in folder_lower or 'about' in folder_lower:
            return 'org_info'
        elif 'example' in folder_lower or 'grant' in folder_lower:
            return 'past_grant'

        # checking filename keywords
        for doc_type, keywords in cls.DOCUMENT_TYPE_KEYWORDS.items():
            for keyword in keywords:
                if keyword in filename_lower or keyword in content_lower[:500]:
                    return doc_type

        # defaulting to past_grant for most historical documents
        return 'past_grant'


class SemanticChunker:
    """
    chunking documents semantically while preserving structure and context
    """

    def __init__(
        self,
        chunk_size: int = 800,
        chunk_overlap: int = 100,
        min_chunk_size: int = 200,
        max_chunk_size: int = 1200,
    ):
        """
        initializing semantic chunker with configurable parameters
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
        self.max_chunk_size = max_chunk_size

        # creating langchain recursive text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=['\n\n## ', '\n\n', '\n', '. ', ' ', ''],
        )

    async def chunk_document(
        self,
        content: str,
        document_metadata: Dict[str, Any],
    ) -> List[Dict[str, Any]]:
        """
        chunking document into semantically meaningful pieces with context
        """
        logger.info(
            "chunking_document",
            filename=document_metadata.get('filename'),
            content_length=len(content),
        )

        # splitting content into chunks
        raw_chunks = self.text_splitter.split_text(content)

        chunks = []
        for idx, chunk_text in enumerate(raw_chunks):
            chunk_length = len(chunk_text)

            # skipping chunks that are too small or too large
            if chunk_length < self.min_chunk_size:
                logger.debug("chunk_too_small", index=idx, length=chunk_length)
                continue

            if chunk_length > self.max_chunk_size:
                # re-splitting large chunks
                logger.debug("chunk_too_large", index=idx, length=chunk_length)
                sub_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=self.max_chunk_size,
                    chunk_overlap=50,
                )
                sub_chunks = sub_splitter.split_text(chunk_text)
                for sub_idx, sub_chunk in enumerate(sub_chunks):
                    chunks.append(self._create_chunk(
                        sub_chunk,
                        document_metadata,
                        f"{idx}.{sub_idx}",
                    ))
            else:
                chunks.append(self._create_chunk(
                    chunk_text,
                    document_metadata,
                    str(idx),
                ))

        logger.info(
            "document_chunked",
            filename=document_metadata.get('filename'),
            chunks_created=len(chunks),
        )

        return chunks

    def _create_chunk(
        self,
        text: str,
        document_metadata: Dict[str, Any],
        chunk_index: str,
    ) -> Dict[str, Any]:
        """
        creating chunk with contextual headers and metadata
        """
        # extracting section name from chunk if present
        section_match = re.search(r'##\s+(.+?)(?:\n|$)', text)
        section_name = section_match.group(1) if section_match else None

        # building contextual header
        context_parts = []
        if document_metadata.get('filename'):
            context_parts.append(f"Document: {document_metadata['filename']}")
        if section_name:
            context_parts.append(f"Section: {section_name}")
        context_parts.append(f"Chunk: {chunk_index}")

        contextual_header = ' | '.join(context_parts)

        # creating chunk with enhanced text including context
        enhanced_text = f"[{contextual_header}]\n\n{text}"

        chunk = {
            'text': enhanced_text,
            'original_text': text,
            'chunk_index': chunk_index,
            'section_name': section_name,
            'char_count': len(text),
            'source_document': document_metadata.get('filename'),
            'document_type': document_metadata.get('document_type'),
            'program': document_metadata.get('program'),
            'funder': document_metadata.get('funder'),
            'year': document_metadata.get('year'),
        }

        return chunk


class DocumentIngestionService:
    """
    orchestrating document loading, metadata extraction, and chunking
    """

    def __init__(self):
        """
        initializing document ingestion service
        """
        self.loader = DocumentLoader()
        self.metadata_extractor = MetadataExtractor()
        self.chunker = SemanticChunker()

    async def scan_directory(self, directory: str) -> List[str]:
        """
        scanning directory recursively for supported document files
        """
        supported_extensions = {'.txt', '.docx', '.pdf'}
        file_paths = []

        directory_path = Path(directory)
        if not directory_path.exists():
            logger.warning("directory_not_found", directory=directory)
            return []

        for file_path in directory_path.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                file_paths.append(str(file_path))

        logger.info("directory_scanned", directory=directory, files_found=len(file_paths))
        return file_paths

    async def process_file(self, file_path: str) -> Dict[str, Any]:
        """
        processing single file through complete ingestion pipeline
        """
        try:
            logger.info("processing_file", file_path=file_path)

            # loading file content
            content, mime_type = await self.loader.load_file(file_path)

            # extracting metadata from filename and content
            path_obj = Path(file_path)
            filename = path_obj.name
            folder_name = path_obj.parent.name

            filename_metadata = self.metadata_extractor.extract_from_filename(filename)
            content_metadata = self.metadata_extractor.extract_from_content(content)

            # merging metadata
            document_metadata = {
                'filename': filename,
                'file_path': file_path,
                'folder_name': folder_name,
                'mime_type': mime_type,
                'file_size_bytes': os.path.getsize(file_path),
                **filename_metadata,
                **content_metadata,
            }

            # inferring document type
            document_type = self.metadata_extractor.infer_document_type(
                filename, content, folder_name
            )
            document_metadata['document_type'] = document_type

            # chunking document
            chunks = await self.chunker.chunk_document(content, document_metadata)

            result = {
                'success': True,
                'file_path': file_path,
                'content': content,
                'metadata': document_metadata,
                'chunks': chunks,
                'chunk_count': len(chunks),
            }

            logger.info(
                "file_processed",
                file_path=file_path,
                chunks=len(chunks),
                document_type=document_type,
            )

            return result

        except Exception as e:
            logger.error("file_processing_error", file_path=file_path, error=str(e))
            return {
                'success': False,
                'file_path': file_path,
                'error': str(e),
            }

    async def process_directory(self, directory: str) -> List[Dict[str, Any]]:
        """
        processing all files in directory through ingestion pipeline
        """
        logger.info("processing_directory", directory=directory)

        # scanning for files
        file_paths = await self.scan_directory(directory)

        if not file_paths:
            logger.warning("no_files_found", directory=directory)
            return []

        # processing files concurrently
        tasks = [self.process_file(file_path) for file_path in file_paths]
        results = await asyncio.gather(*tasks)

        successful = sum(1 for r in results if r.get('success'))
        failed = len(results) - successful

        logger.info(
            "directory_processed",
            directory=directory,
            total=len(results),
            successful=successful,
            failed=failed,
        )

        return results


# creating singleton instance
_ingestion_service: Optional[DocumentIngestionService] = None


def get_ingestion_service() -> DocumentIngestionService:
    """
    getting singleton document ingestion service
    """
    global _ingestion_service
    if _ingestion_service is None:
        _ingestion_service = DocumentIngestionService()
    return _ingestion_service
